import sys
import io
from docx import Document
import json

# 设置控制台输出为 UTF-8
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8', errors='replace')

# 读取 Word 文档
doc = Document(r'C:\Users\1\Downloads\磷酸铁锂行业综合洞察报告_2026Q1.docx')

# 提取所有内容
content = {
    'paragraphs': [],
    'tables': []
}

# 提取段落
for i, para in enumerate(doc.paragraphs):
    if para.text.strip():
        content['paragraphs'].append({
            'index': i,
            'text': para.text,
            'style': para.style.name if para.style else None
        })

# 提取表格
for i, table in enumerate(doc.tables):
    table_data = []
    for row in table.rows:
        row_data = []
        for cell in row.cells:
            row_data.append(cell.text)
        table_data.append(row_data)
    content['tables'].append({
        'index': i,
        'data': table_data
    })

# 打印内容
print("=" * 100)
print("文档内容提取")
print("=" * 100)

print(f"\n共找到 {len(content['paragraphs'])} 个段落")
print(f"共找到 {len(content['tables'])} 个表格\n")

print("=" * 100)
print("段落内容:")
print("=" * 100)
for para in content['paragraphs']:
    print(f"\n[{para['index']}] ({para['style']})")
    print(para['text'])

print("\n" + "=" * 100)
print("表格内容:")
print("=" * 100)
for table in content['tables']:
    print(f"\n表格 {table['index'] + 1}:")
    print("-" * 100)
    for i, row in enumerate(table['data']):
        print(f"行 {i + 1}: {row}")

# 保存为 JSON 文件以便后续使用
with open('report_content.json', 'w', encoding='utf-8') as f:
    json.dump(content, f, ensure_ascii=False, indent=2)

print("\n" + "=" * 100)
print("内容已保存到 report_content.json")
print("=" * 100)
